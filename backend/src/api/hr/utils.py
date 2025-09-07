from datetime import datetime
from decimal import Decimal, InvalidOperation
import re

from fastapi import UploadFile


def format_datetime(dt: datetime) -> str:
    """Форматирует дату и время в ISO-формат."""
    return dt.isoformat() if dt else None


def _clean_text(s: str) -> str:
    return (s or "").replace("\u00A0", " ").strip()


def to_decimal(value) -> Decimal:
    """Приводит число вида '120 000 руб.' или '15,5' к Decimal(120000 / 15.5)."""
    if value is None:
        return Decimal(0)
    try:
        s = str(value)
        s = s.replace("\u00A0", " ")  # неразрывные пробелы
        s = s.strip()
        # оставляем только цифры, точки, запятые и знак минус
        s = re.sub(r"[^\d,.\-]", "", s)
        s = s.replace(" ", "").replace(",", ".")
        if not s or s in {".", "-", "-.", ".-"}:
            return Decimal(0)
        return Decimal(s)
    except (InvalidOperation, Exception):
        return Decimal(0)


def parse_experience(val) -> int:
    if not val:
        return 0
    s = str(val).lower().strip()
    if any(k in s for k in ["нет опыта", "без опыта", "не требуется", "не определ", "n/a", "na"]):
        return 0
    s = s.replace("—", "-").replace("–", "-").replace("+", "")
    nums = re.findall(r"\d+", s)
    if not nums:
        return 0
    years = [int(n) for n in nums]
    if "до" in s and years:
        return years[-1]
    if "от" in s and years:
        return years[0]
    if "более" in s or "не менее" in s:
        return years[0]
    # диапазон '1-3' -> берём минимум
    return min(years)


def parse_bool_ru(val) -> bool:
    s = str(val or "").strip().lower()
    return s in {"да", "есть", "true", "1", "y", "yes"}


def parse_vacancy_docx(upload: UploadFile) -> dict:
    try:
        from docx import Document
    except ModuleNotFoundError:
        raise ValueError("Зависимость python-docx не установлена. Установи пакет: pip install python-docx")

    if not upload or not upload.filename or not upload.filename.lower().endswith(".docx"):
        raise ValueError("Ожидается DOCX-файл.")

    try:
        upload.file.seek(0)
        doc = Document(upload.file)
    except Exception:
        raise ValueError("Не удалось прочитать DOCX. Проверь формат файла.")

    data: dict[str, str] = {}

    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) < 2:
                continue
            key = _clean_text(row.cells[0].text)
            val = _clean_text(row.cells[1].text)
            if key:
                data[key] = val
        if data:
            break

    if not data:
        paras = [_clean_text(p.text) for p in doc.paragraphs if _clean_text(p.text)]
        for i in range(0, len(paras) - 1, 2):
            key = paras[i]
            val = paras[i + 1]
            if key and val:
                data[key] = val

    if not data:
        raise ValueError("В документе не найдены пары «Наименование поля / Значение».")
    return data


def vacancy_to_txt(file, as_text=False):
    """
    Маппинг русских полей -> внутренние поля Vacancy.
    Числа/проценты оставляем строкой — в БД приводим через to_decimal().
    """
    offer_type_map = {
        "Постоянно": "TK",
        "ГПХ": "GPH",
        "ИП": "IP",
    }
    busy_type_map = {
        "Полная занятость": "allTime",
        "Проектная работа": "projectTime",
    }
    vacancy_status_map = {
        "Открыта": "active",
        "Найден кандидат": "closed",
        "Одобрена": "active",
        "Приостановлена": "stopped",
    }

    vacancy_data = {
        "name": file.get("Название", ""),
        "region": file.get("Регион", ""),
        "offerType": offer_type_map.get(file.get("Тип трудового", ""), "TK"),
        "busyType": busy_type_map.get(file.get("Тип занятости", ""), "allTime"),
        "graph": file.get("Текст график работы", ""),
        "salaryMin": file.get("Оклад мин. (руб/мес)", None),
        "salaryMax": file.get("Оклад макс. (руб/мес)", None),
        "annualBonus": file.get("Годовая премия (%)", None),
        "description": file.get("Обязанности (для публикации)", ""),
        "promt": file.get("Требования (для публикации)", ""),
        "exp": parse_experience(file.get("Требуемый опыт работы", "")),
        "degree": str(file.get("Уровень образования", "") or "").strip().lower().startswith("высше"),
        # Доп. поля для модели Vacancy
        "city": file.get("Город", ""),
        "address": file.get("Адрес", ""),
        "status": vacancy_status_map.get(file.get("Статус", ""), "active"),
        "bonusType": file.get("Тип премирования. Описание", ""),
        "specialSoftware": file.get("Знание специальных программ", ""),
        "computerSkills": file.get("Навыки работы на компьютере", ""),
        "foreignLanguages": file.get("Знание иностранных языков", ""),
        "languageLevel": file.get("Уровень владения языка", ""),
        "businessTrips": parse_bool_ru(file.get("Наличие командировок", "")),
        "date": format_datetime(datetime.now()),
    }

    if as_text:
        result = ["=== Вакансия ==="]
        for key, value in vacancy_data.items():
            if value not in (None, "", []):
                if key in ["description", "promt"]:
                    result.append(f"{key}:")
                    for line in str(value).split(";"):
                        line = line.strip()
                        if line:
                            result.append(f"  - {line}")
                else:
                    result.append(f"{key}: {value}")
        return "\n".join(result)

    return vacancy_data